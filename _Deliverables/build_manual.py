"""Build Instruction_Manual.docx for JabberJuicy."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

IMG = os.path.join(os.path.dirname(__file__), "Instruction_Manuel")
OUT = os.path.join(os.path.dirname(__file__), "Instruction_Manual.docx")

ORANGE = RGBColor(0xC0, 0x50, 0x00)
GREEN  = RGBColor(0x2F, 0x7A, 0x3C)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def heading(text, level=1, color=ORANGE):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    return p

def body(text, italic=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_font(run, 11, italic=italic)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_font(r1, 11, bold=True)
        r2 = p.add_run(text)
        set_font(r2, 11)
    else:
        r = p.add_run(text)
        set_font(r, 11)

def img(filename, width=Inches(5.5), caption=None):
    path = os.path.join(IMG, filename)
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=width)
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        for run in cp.runs:
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def add_toc():
    """Insert a Word TOC field that populates when opened in Word."""
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = toc_heading.add_run("Table of Contents")
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = ORANGE

    p = doc.add_paragraph()
    run = p.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    # Placeholder text shown before Word updates the field
    noUpdate = OxmlElement('w:r')
    noUpdateT = OxmlElement('w:t')
    noUpdateT.text = '[Right-click and select "Update Field" to generate the table of contents]'
    noUpdate.append(noUpdateT)
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_sep)
    run._r.append(noUpdate)
    run._r.append(fldChar_end)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

def rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C05000')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("JabberJuicy")
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = ORANGE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("User Instruction Manual")
r2.font.size = Pt(18)
r2.font.color.rgb = GREEN

doc.add_paragraph()
quote = doc.add_paragraph()
quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
rq = quote.add_run('"All mimsy were the borogoves, and the mome raths outgrabe."')
rq.font.size = Pt(11)
rq.font.italic = True
rq.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
team = doc.add_paragraph()
team.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = team.add_run("Echo Team  ·  DASC Data Management  ·  Spring 2026\nGrayson Bandy  ·  Warren Ross  ·  Adrian Munoz  ·  Andrew Cox")
rt.font.size = Pt(11)
rt.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
add_toc()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
heading("1. Introduction")
rule()
body(
    "JabberJuicy is an online ordering application for a fictional juice bar chain inspired by "
    "Lewis Carroll's poem Jabberwocky. Customers can browse a whimsical menu of juices and "
    "smoothies, build an order, check out, and track their order history — all from any web browser."
)
body(
    "The application is live at:"
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
r = p.add_run("https://jabberjuicy-production.up.railway.app")
r.font.size = Pt(11)
r.font.bold = True
r.font.color.rgb = GREEN

body("")
body("This manual walks through every page of the app, explains each feature, and provides "
     "step-by-step instructions for common tasks.")

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 2. GETTING STARTED
# ══════════════════════════════════════════════════════════════════════════════
heading("2. Getting Started")
rule()

heading("2.1  Creating an Account", level=2, color=GREEN)
body(
    "New users must register before placing an order. To create an account:"
)
bullet("Open the app in any modern web browser.")
bullet("Click Register in the top navigation bar.")
bullet("Fill in the required fields:", bold_prefix="Fill in required fields:  ")
p2 = doc.add_paragraph(style="List Bullet 2")
p2.paragraph_format.space_after = Pt(2)
p2.add_run("First Name, Last Name, Username, Password (all required)").font.size = Pt(11)
p3 = doc.add_paragraph(style="List Bullet 2")
p3.paragraph_format.space_after = Pt(2)
p3.add_run("Email, Phone, Address, City, State, Zip Code (optional but recommended)").font.size = Pt(11)
bullet("Click Create Account. You will be logged in automatically.")

doc.add_paragraph()
img("register.png", caption="Figure 1 — Register page (/register)")

doc.add_paragraph()
heading("2.2  Logging In", level=2, color=GREEN)
body(
    "If you already have an account:"
)
bullet("Click Login in the top navigation bar.")
bullet("Enter your Username and Password.")
bullet("Click Login. You will be taken to your personal dashboard.")
body(
    "If you enter incorrect credentials the page will display an error — check your username "
    "and password and try again.",
    italic=True
)
doc.add_paragraph()
img("login.png", caption="Figure 2 — Login page (/login)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. DASHBOARD (HOME)
# ══════════════════════════════════════════════════════════════════════════════
heading("3. Dashboard")
rule()
body(
    "After logging in you land on the Dashboard. It gives a quick overview of everything "
    "the app offers and shows your current JabberWonk Points balance."
)

heading("Dashboard cards", level=2, color=GREEN)
bullet("New Order — Start building a new order right away.")
bullet("View Menu — Browse the full drink catalog without committing to an order.")
bullet("Order History — Review every past order and its status.")
bullet("JabberWonk Points — Displays your current point balance, the earn rate, and the redemption rate.")

body(
    "The top navigation bar also shows a green Pickup Ready badge whenever you have a pending "
    "order waiting to be picked up.",
    italic=False
)
doc.add_paragraph()
img("home.png", caption="Figure 3 — Dashboard (/home)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. MENU
# ══════════════════════════════════════════════════════════════════════════════
heading("4. Menu")
rule()
body(
    "The Menu page is publicly accessible — no login required. It displays the full catalog "
    "of available drinks organized in a card grid."
)
bullet("Each card shows the drink name, a short description, category badge (Juice / Smoothie / Special), and price.")
bullet("Use the menu to decide what to order before starting a new order.")
body(
    "To begin ordering, log in and click New Order in the navigation bar or on the Dashboard.",
    italic=True
)
doc.add_paragraph()
img("menu.png", caption="Figure 4 — Menu page (/menu)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. PLACING AN ORDER
# ══════════════════════════════════════════════════════════════════════════════
heading("5. Placing an Order")
rule()

heading("5.1  Build Your Order", level=2, color=GREEN)
body(
    "The Build Your Order page is split into two panels:"
)
bullet("Add Items (left) — Select a drink from the dropdown and enter a quantity, then click Add to Order.")
bullet("Your Cart (right) — Shows every item added so far, with quantity, unit price, subtotal per item, and the running order total.")
bullet("Remove — Click the Remove button next to any line item to delete it from the cart.")
bullet("Clear Cart — Removes all items and starts fresh.")
bullet("Proceed to Checkout — When you are happy with your cart, click this button to move to checkout.")
doc.add_paragraph()
img("order.png", caption="Figure 5 — Build Your Order page (/order)")

doc.add_paragraph()
heading("5.2  Checkout", level=2, color=GREEN)
body(
    "The Checkout page lets you finalize the order before submitting it."
)
bullet("Order Summary (left) — A read-only recap of your cart.")
bullet("← Edit Order — Takes you back to the cart if you need to make changes.")
bullet("Pick-up Location — Choose which JabberJuicy location you will pick up from.")
bullet("Payment Method — Select Cash, Credit Card, Debit Card, or JabberWonk Points.")
p_tip = doc.add_paragraph(style="List Bullet")
p_tip.paragraph_format.space_after = Pt(3)
rt = p_tip.add_run("JabberWonk Points tip:  ")
rt.font.size = Pt(11)
rt.font.bold = True
rb = p_tip.add_run(
    "If your balance is enough to cover the order total the app will show "
    "a green callout confirming you can pay with points. If your balance is insufficient "
    "the option will be unavailable."
)
rb.font.size = Pt(11)
bullet("Special Notes — Optional field for allergies, substitutions, or special requests.")
bullet("Complete Transaction — Submits the order. You will be taken to the receipt page.")
doc.add_paragraph()
img("checkout.png", caption="Figure 6 — Checkout page (/checkout)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. RECEIPT
# ══════════════════════════════════════════════════════════════════════════════
heading("6. Order Receipt")
rule()
body(
    "After completing checkout you are shown an order receipt confirming all the details."
)
bullet("Order number, date/time, location, and payment method are displayed.")
bullet("A line-item table shows everything you ordered with the total paid.")
bullet("If you paid with cash or card, a gold banner shows how many JabberWonk Points you earned.")
bullet("If you paid with JabberWonk Points, the banner shows how many points were redeemed.")

body("From the receipt you can:")
bullet("Confirm Pickup — Jump straight to the pickup page for this order.")
bullet("Place Another Order — Start a new cart.")
bullet("View Order History — See all past orders.")
bullet("Main Menu — Return to the Dashboard.")
doc.add_paragraph()
img("receipt.png", caption="Figure 7 — Order Confirmed receipt (/receipt)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. PICKUP
# ══════════════════════════════════════════════════════════════════════════════
heading("7. Confirming Pickup")
rule()
body(
    "Orders stay in Pending status until you physically pick them up and confirm in the app. "
    "The green Pickup Ready badge in the navigation bar tells you when you have orders waiting."
)

heading("7.1  Pickup Confirmation Page", level=2, color=GREEN)
body(
    "Click the Pickup Ready badge (or navigate to /pickup) to reach the pickup page. "
    "If you have more than one pending order you will first see a list to choose from."
)
body(
    "The Confirm Pickup — Order #XX page shows:"
)
bullet("Order date, location, and payment method.")
bullet("A full item table with quantities, prices, and total.")
bullet("Confirm Pickup button — marks the order Completed and takes you to the success page.")
bullet("Cancel Order button — cancels the order. If the order was paid with JabberWonk Points, "
       "the full points are automatically refunded to your balance.")
doc.add_paragraph()
img("pickup-confirm.png", caption="Figure 8 — Confirm Pickup page (/pickup/{id})")

doc.add_paragraph()
heading("7.2  Pickup Success", level=2, color=GREEN)
body(
    "After confirming pickup you see the Order Picked Up! success page. "
    "It includes a Jabberwocky-themed quote matched to the drinks in your order. "
    "From here you can place another order, view your history, or return home."
)
doc.add_paragraph()
img("pickup-success.png", caption="Figure 9 — Pickup success page (/pickup/{id}/success)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. ORDER HISTORY
# ══════════════════════════════════════════════════════════════════════════════
heading("8. Order History")
rule()
body(
    "The Order History page lists every order you have ever placed, sorted newest first."
)
bullet("Each row shows: Order #, Date, Location, Items ordered, Total, and Status.")
bullet("Status badges are color-coded: green = Completed, yellow = Pending, red = Cancelled.")
bullet("Click New Order (top right) to start a fresh cart at any time.")
doc.add_paragraph()
img("history.png", caption="Figure 10 — My Order History (/history)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. JABBERWONK POINTS
# ══════════════════════════════════════════════════════════════════════════════
heading("9. JabberWonk Points")
rule()
body(
    "JabberWonk Points are JabberJuicy's loyalty reward system. Every qualifying order "
    "earns points that can be redeemed for free drinks on future orders."
)

heading("Earning Points", level=2, color=GREEN)
bullet("50 points are awarded as a flat visit bonus on every order paid with cash or card.")
bullet("10 additional points are earned for every whole dollar spent (e.g. a $7.50 order earns 70 pts from spend).")
bullet("Orders paid with JabberWonk Points do not earn additional points (no double-dipping).")

heading("Redeeming Points", level=2, color=GREEN)
bullet("100 points = $1.00 off your order.")
bullet("At checkout, select JabberWonk Points as your payment method.")
bullet("The app will show your current balance and whether you have enough to cover the order.")
bullet("Points are deducted from your balance when the order is placed.")

heading("Refunds on Cancellation", level=2, color=GREEN)
bullet(
    "If you cancel an order that was paid with JabberWonk Points, "
    "the full points are automatically refunded to your account. "
    "The refund is logged in your points history."
)

heading("Checking Your Balance", level=2, color=GREEN)
bullet("Your current balance is always shown on the Dashboard.")
bullet("The navigation bar displays your running point total when you are logged in.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. ACCOUNT SETTINGS
# ══════════════════════════════════════════════════════════════════════════════
heading("10. Account Settings")
rule()
body(
    "You can update your username from the Account Settings page."
)
bullet("Click your username in the top navigation bar to reveal a dropdown.")
bullet("Select Account Settings.")
bullet("Enter your new username and click Save Username.")
bullet("If the new username is already taken, an error will appear — choose a different one.")
bullet("On success you will see a confirmation page and be logged in under the new username.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. NAVIGATION REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
heading("11. Quick Navigation Reference")
rule()
body("The table below lists every page and whether a login is required to access it.")
doc.add_paragraph()

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
for cell, text in zip(hdr, ["Page", "URL", "Login Required"]):
    cell.text = text
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(11)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

rows = [
    ("Landing / Hero",           "/",                    "No"),
    ("Register",                 "/register",            "No"),
    ("Login",                    "/login",               "No"),
    ("Menu",                     "/menu",                "No"),
    ("Dashboard",                "/home",                "Yes"),
    ("Build Your Order",         "/order",               "Yes"),
    ("Checkout",                 "/checkout",            "Yes"),
    ("Order Receipt",            "/receipt",             "Yes"),
    ("Order History",            "/history",             "Yes"),
    ("Account Settings",         "/account",             "Yes"),
    ("Select Pending Pickup",    "/pickup",              "Yes"),
    ("Confirm / Cancel Pickup",  "/pickup/{id}",         "Yes"),
    ("Pickup Success",           "/pickup/{id}/success", "Yes"),
]
for page, url, login in rows:
    row = table.add_row().cells
    row[0].text = page
    row[1].text = url
    row[2].text = login
    for c in row:
        c.paragraphs[0].paragraph_format.space_after = Pt(2)
        for run in c.paragraphs[0].runs:
            run.font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUT)
print(f"Saved: {OUT}")
